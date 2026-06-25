"""Document parsing factory for RAG system.

Provides unified interface for parsing various document formats including
TXT, MD, HTML, PDF, DOCX, EPUB, and MOBI files.
"""

import asyncio
import concurrent.futures
import re
import subprocess  # nosec B404 - required for poppler-utils check
from pathlib import Path

from bot.utils.logging import get_logger

logger = get_logger(__name__)


class DocumentParserFactory:
    """Factory for creating document parsers based on file type."""

    def __init__(self) -> None:
        self._supported_extensions = {
            ".txt",
            ".md",
            ".markdown",
            ".html",
            ".htm",
            ".pdf",
            ".docx",
            ".epub",
            ".mobi",
        }

    def get_supported_extensions(self) -> set[str]:
        """Get set of supported file extensions."""
        return self._supported_extensions.copy()

    def get_parser(self, file_path: str | Path) -> str | None:
        """Get parser type for given file path."""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()

        if extension in self._supported_extensions:
            return extension[1:]  # Remove the dot
        return None

    async def parse_document(self, file_path: str | Path) -> tuple[str, dict]:
        """Parse document and return content and metadata.

        Args:
            file_path: Path to the document to parse

        Returns:
            Tuple of (content, metadata)

        Raises:
            ValueError: If file type is not supported
            FileNotFoundError: If file does not exist

        """
        file_path = Path(file_path)

        if not file_path.exists():
            msg = f"File not found: {file_path}"
            raise FileNotFoundError(msg)

        extension = file_path.suffix.lower()

        if extension not in self._supported_extensions:
            msg = f"Unsupported file type: {extension}"
            raise ValueError(msg)

        logger.debug(f"Parsing document: {file_path} (type: {extension})")

        try:
            if extension in {".txt", ".md", ".markdown"}:
                return await self._parse_text_file(file_path, extension)
            if extension in {".html", ".htm"}:
                return await self._parse_html_file(file_path)
            if extension == ".pdf":
                return await self._parse_pdf_file(file_path)
            if extension == ".docx":
                return await self._parse_docx_file(file_path)
            if extension == ".epub":
                return await self._parse_epub_file(file_path)
            if extension == ".mobi":
                return await self._parse_mobi_file(file_path)
            msg = f"Parser not implemented for: {extension}"
            raise ValueError(msg)

        except Exception as e:
            logger.exception(f"Failed to parse {file_path}: {e}")
            raise

    async def _parse_text_file(self, file_path: Path, extension: str) -> tuple[str, dict]:
        """Parse plain text or markdown files."""
        try:
            with open(file_path, encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            # Fallback to latin-1 encoding
            with open(file_path, encoding="latin-1") as f:
                content = f.read()

        metadata = {
            "parser_type": "text",
            "file_type": extension.removeprefix("."),
            "char_count": len(content),
            "line_count": len(content.splitlines()),
            "file_size": file_path.stat().st_size,
        }

        # Add markdown-specific metadata
        if extension in {".md", ".markdown"}:
            metadata["parser_type"] = "markdown"
            headers = re.findall(r"^#+\s+(.+)$", content, re.MULTILINE)
            metadata["header_count"] = len(headers)
            if headers:
                metadata["first_header"] = headers[0]

        return content.strip(), metadata

    async def _parse_html_file(self, file_path: Path) -> tuple[str, dict]:
        """Parse HTML files."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            logger.warning("BeautifulSoup not available, falling back to basic HTML parsing")
            return await self._parse_text_file(file_path, ".html")

        try:
            with open(file_path, encoding="utf-8") as f:
                html_content = f.read()
        except UnicodeDecodeError:
            with open(file_path, encoding="latin-1") as f:
                html_content = f.read()

        soup = BeautifulSoup(html_content, "html.parser")

        # Extract title
        title = soup.find("title")
        title_text = title.get_text().strip() if title else ""

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        # Get text content
        content = soup.get_text()

        # Clean up whitespace
        lines = (line.strip() for line in content.splitlines())
        content = "\n".join(line for line in lines if line)

        metadata = {
            "parser_type": "html",
            "file_type": "html",
            "html_title": title_text,
            "char_count": len(content),
            "line_count": len(content.splitlines()),
            "file_size": file_path.stat().st_size,
        }

        return content.strip(), metadata

    async def _parse_pdf_file(self, file_path: Path) -> tuple[str, dict]:
        """Parse PDF files with OCR fallback."""
        content = ""
        metadata = {
            "parser_type": "pdf",
            "file_type": "pdf",
            "file_size": file_path.stat().st_size,
        }

        # Try PyPDF2 first (run in thread pool to avoid blocking)
        try:
            loop = asyncio.get_event_loop()

            # Run PDF reading in thread pool to prevent blocking
            content = await loop.run_in_executor(None, self._extract_pdf_with_pypdf2, file_path)

            if content:
                # Get page count for metadata
                try:
                    page_count = await loop.run_in_executor(None, self._get_pdf_page_count, file_path)
                    metadata["page_count"] = page_count
                except (ImportError, AttributeError, RuntimeError, OSError):
                    metadata["page_count"] = "unknown"

                metadata["extraction_method"] = "pypdf2"

        except ImportError:
            logger.warning("PyPDF2 not available, trying OCR fallback")
            content = ""
        except (ImportError, RuntimeError, OSError, ValueError, AttributeError) as e:
            logger.warning(f"PyPDF2 extraction failed: {e}, trying OCR fallback")
            content = ""

        # If no content extracted, try OCR fallback
        if not content.strip():
            try:
                content = await self._ocr_pdf_fallback(file_path)
                metadata["extraction_method"] = "ocr"
            except (ImportError, RuntimeError, OSError, ValueError, subprocess.CalledProcessError, FileNotFoundError) as e:
                logger.exception(f"OCR fallback failed for {file_path}: {e}")
                msg = f"Could not extract text from PDF: {e}"
                raise ValueError(msg)

        metadata["char_count"] = len(content)
        metadata["line_count"] = len(content.splitlines())

        return content.strip(), metadata

    def _extract_pdf_with_pypdf2(self, file_path: Path) -> str:
        """Extract text from PDF using PyPDF2. Runs in thread pool."""
        try:
            import PyPDF2

            content = ""
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)

                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            content += f"[Page {page_num + 1}]\n{page_text}\n\n"
                    except (AttributeError, RuntimeError, ValueError) as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue

            return content

        except (ImportError, RuntimeError, OSError, ValueError, AttributeError) as e:
            logger.exception(f"PyPDF2 extraction failed: {e}")
            return ""

    def _get_pdf_page_count(self, file_path: Path) -> int:
        """Get PDF page count. Runs in thread pool."""
        try:
            import PyPDF2

            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                return len(reader.pages)

        except (ImportError, RuntimeError, OSError, ValueError, AttributeError) as e:
            logger.warning(f"Failed to get page count: {e}")
            return 0

    async def _ocr_pdf_fallback(self, file_path: Path) -> str:
        """OCR fallback for PDF files using pdf2image and pytesseract."""
        try:
            from pdf2image import convert_from_path
        except ImportError as e:
            msg = f"OCR dependencies not available: {e}. Install with: pip install pdf2image"
            raise ImportError(msg)

        # Check if poppler-utils is available
        try:
            subprocess.run(["pdftoppm", "-h"], capture_output=True, check=True)  # nosec B603, B607
        except (subprocess.CalledProcessError, FileNotFoundError):
            msg = "poppler-utils not found. Install with: sudo apt-get install poppler-utils"
            raise ImportError(msg)

        logger.info(f"Using OCR to extract text from {file_path}")

        # Convert PDF to images (this can also be slow, so run in thread pool)
        loop = asyncio.get_event_loop()
        try:
            images = await loop.run_in_executor(None, lambda: convert_from_path(str(file_path)))
        except (RuntimeError, OSError, ValueError, ImportError, subprocess.CalledProcessError, FileNotFoundError) as e:
            logger.exception(f"Failed to convert PDF to images: {e}")
            msg = f"Could not convert PDF to images: {e}"
            raise ValueError(msg)

        logger.info(f"Processing {len(images)} pages with OCR")

        # Process pages with OCR in thread pool to avoid blocking
        content = ""
        with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
            # Create tasks for OCR processing
            ocr_tasks = []
            for page_num, image in enumerate(images):
                future = loop.run_in_executor(executor, self._ocr_single_page, image, page_num + 1)
                ocr_tasks.append(future)

            # Wait for all OCR tasks with timeout
            try:
                results = await asyncio.wait_for(
                    asyncio.gather(*ocr_tasks, return_exceptions=True),
                    timeout=300,  # 5 minute timeout for entire OCR process
                )

                # Collect results
                for page_num, result in enumerate(results):
                    if isinstance(result, Exception):
                        logger.warning(f"OCR failed for page {page_num + 1}: {result}")
                    elif result:
                        content += result

            except TimeoutError:
                logger.exception(f"OCR processing timed out for {file_path}")
                msg = "OCR processing timed out - PDF may be too large or complex"
                raise ValueError(msg)

        return content

    def _ocr_single_page(self, image, page_num: int) -> str:
        """OCR a single page image. Runs in thread pool."""
        try:
            import pytesseract

            # Set timeout for individual page OCR
            page_text = pytesseract.image_to_string(image, timeout=60)  # 1 minute per page

            if page_text.strip():
                return f"[Page {page_num}]\n{page_text}\n\n"
            return ""

        except (RuntimeError, OSError, ValueError, ImportError) as e:
            logger.warning(f"OCR failed for page {page_num}: {e}")
            return ""

    async def _parse_docx_file(self, file_path: Path) -> tuple[str, dict]:
        """Parse DOCX files."""
        try:
            from docx import Document
        except ImportError:
            msg = "python-docx not available. Install with: pip install python-docx"
            raise ImportError(msg)

        doc = Document(str(file_path))

        content = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content += paragraph.text + "\n"

        metadata = {
            "parser_type": "docx",
            "file_type": "docx",
            "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
            "char_count": len(content),
            "line_count": len(content.splitlines()),
            "file_size": file_path.stat().st_size,
        }

        return content.strip(), metadata

    async def _parse_epub_file(self, file_path: Path) -> tuple[str, dict]:
        """Parse EPUB files."""
        try:
            import ebooklib
            from bs4 import BeautifulSoup
            from ebooklib import epub
        except ImportError:
            msg = "ebooklib and beautifulsoup4 not available. Install with: pip install ebooklib beautifulsoup4"
            raise ImportError(msg)

        book = epub.read_epub(str(file_path))

        content = ""
        chapter_count = 0

        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                chapter_count += 1
                soup = BeautifulSoup(item.get_content(), "html.parser")

                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()

                chapter_text = soup.get_text()

                # Clean up whitespace
                lines = (line.strip() for line in chapter_text.splitlines())
                chapter_text = "\n".join(line for line in lines if line)

                if chapter_text.strip():
                    content += f"[Chapter {chapter_count}]\n{chapter_text}\n\n"

        # Get book metadata
        title = book.get_metadata("DC", "title")
        author = book.get_metadata("DC", "creator")

        metadata = {
            "parser_type": "epub",
            "file_type": "epub",
            "chapter_count": chapter_count,
            "char_count": len(content),
            "line_count": len(content.splitlines()),
            "file_size": file_path.stat().st_size,
        }

        if title:
            metadata["book_title"] = title[0][0] if title else ""
        if author:
            metadata["book_author"] = author[0][0] if author else ""

        return content.strip(), metadata

    async def _parse_mobi_file(self, file_path: Path) -> tuple[str, dict]:
        """Parse MOBI files (basic support)."""
        # MOBI parsing is complex and requires specialized libraries
        # For now, we'll provide basic support by trying to extract readable text
        logger.warning(f"MOBI parsing has limited support for {file_path}")

        try:
            # Try to read as binary and extract readable text
            with open(file_path, "rb") as f:
                binary_content = f.read()

            # Simple heuristic to extract readable text
            # This is very basic and may not work well for all MOBI files
            text_content = ""
            for chunk in binary_content.split(b"\x00"):
                try:
                    decoded = chunk.decode("utf-8", errors="ignore")
                    if len(decoded) > 20 and decoded.isprintable():
                        text_content += decoded + " "
                except (UnicodeDecodeError, ValueError, AttributeError) as exc:
                    logger.debug(f"Failed to decode MOBI chunk: {exc}")
                    continue

            # Clean up the extracted text
            content = re.sub(r"\s+", " ", text_content).strip()

            metadata = {
                "parser_type": "mobi",
                "file_type": "mobi",
                "char_count": len(content),
                "line_count": len(content.splitlines()),
                "file_size": file_path.stat().st_size,
                "extraction_method": "basic_binary",
            }

            return content, metadata

        except (RuntimeError, OSError, ValueError, UnicodeDecodeError, AttributeError) as e:
            msg = f"Failed to parse MOBI file {file_path}: {e}"
            raise ValueError(msg)


# Global factory instance
document_parser_factory = DocumentParserFactory()


# Utility functions for backward compatibility
def get_supported_extensions() -> set[str]:
    """Get supported file extensions."""
    return document_parser_factory.get_supported_extensions()


def get_parser(file_path: str | Path) -> str | None:
    """Get parser type for file."""
    return document_parser_factory.get_parser(file_path)


async def parse_document(file_path: str | Path) -> tuple[str, dict]:
    """Parse document and return content and metadata."""
    return await document_parser_factory.parse_document(file_path)
